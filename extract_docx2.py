import sys
import io
from docx import Document

# Fix Unicode output on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

docx_path = r"C:\Users\swapnonil.mukherjee\projects\vivli-chatbot\Functional_Requirements_Form_Validation_Rules_Professional.docx"

doc = Document(docx_path)

print("=" * 100)
print("FORM VALIDATION - FUNCTIONAL REQUIREMENTS (DEV NOTES & PSEUDO CODE)")
print("=" * 100)
print()

# Print paragraphs
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n" + "=" * 100)
print("KEY TABLES WITH PSEUDO CODE EXAMPLES")
print("=" * 100)
print()

# Print tables focusing on dev notes
for table_num, table in enumerate(doc.tables, 1):
    print(f"\n{'='*100}")
    print(f"Table {table_num}:")
    print(f"{'='*100}")
    for row_num, row in enumerate(table.rows, 1):
        row_data = [cell.text.strip() for cell in row.cells]
        if any(row_data):  # Only print non-empty rows
            print(f"\nRow {row_num}:")
            for col_num, cell_text in enumerate(row_data, 1):
                if cell_text and len(cell_text) > 0:
                    print(f"\n  Column {col_num}:")
                    print(f"  {'-'*95}")
                    print(f"  {cell_text}")
