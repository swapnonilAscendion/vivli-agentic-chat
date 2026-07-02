from docx import Document

docx_path = r"C:\Users\swapnonil.mukherjee\projects\vivli-chatbot\Vivli_Functional_Requirements_Platform_Agentic Chat_Professional.docx"

doc = Document(docx_path)

print("=" * 80)
print("FUNCTIONAL REQUIREMENTS DOCUMENT")
print("=" * 80)
print()

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n" + "=" * 80)
print("TABLES IN DOCUMENT")
print("=" * 80)
print()

for table_num, table in enumerate(doc.tables, 1):
    print(f"\nTable {table_num}:")
    print("-" * 80)
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_data))
    print("-" * 80)
