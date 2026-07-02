import sys
import io
from docx import Document

# Fix Unicode output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

docx_path = r"C:\Users\swapnonil.mukherjee\projects\vivli-chatbot\Vivli_Functional_Requirements_Platform_Agentic Chat_Professional.docx"

doc = Document(docx_path)

print("=" * 120)
print("QUERY INPUT SECTION - FROM FRD")
print("=" * 120)
print()

# Flag to track when we're in the Query Input section
in_query_input = False
section_count = 0

for para in doc.paragraphs:
    text = para.text.strip()

    if "Query Input" in text:
        in_query_input = True
        print(text)
        print()
        continue

    if in_query_input:
        # Stop when we hit the next major section
        if text.startswith("Response Delivery"):
            break

        if text:
            print(text)

print("\n" + "=" * 120)
print("QUERY INPUT TABLES - DEV NOTES")
print("=" * 120)
print()

# Extract tables that are part of Query Input section
for table_num, table in enumerate(doc.tables, 1):
    # Try to identify which table by looking at first cell
    first_cell_text = table.rows[0].cells[0].text.strip() if table.rows else ""

    # Look for tables with "Intent Classification" or validation-related content
    if ("Intent" in first_cell_text or "intent" in first_cell_text or
        "Checks" in first_cell_text or any("check" in cell.text.lower()
        for row in table.rows for cell in row.cells)):

        print(f"\n{'='*120}")
        print(f"TABLE: {first_cell_text}")
        print(f"{'='*120}\n")

        for row_num, row in enumerate(table.rows, 1):
            row_data = [cell.text.strip() for cell in row.cells]

            if any(row_data):
                if row_num == 1:
                    print("HEADERS:")
                    for col_num, cell_text in enumerate(row_data, 1):
                        print(f"  Col {col_num}: {cell_text}")
                    print("-" * 120)
                else:
                    print(f"\nRow {row_num}:")
                    for col_num, cell_text in enumerate(row_data, 1):
                        if cell_text:
                            print(f"  Column {col_num}:")
                            print(f"  {cell_text}")
