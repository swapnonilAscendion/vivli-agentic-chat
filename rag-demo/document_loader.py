import os
import logging
from pathlib import Path
from typing import List, Dict, Optional, Callable
import json
import csv

try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    PdfReader = None

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    Document = None

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False
    openpyxl = None

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load documents from various file formats"""

    def __init__(self, base_path: str = None):
        """
        Initialize document loader.

        Args:
            base_path: Base directory to load documents from
        """
        self.base_path = base_path or "../resources/organized-data"

    def load_text_file(self, file_path: str) -> str:
        """Load text from a file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            return ""

    def load_markdown_files(self, directory: str) -> List[Dict]:
        """
        Load all markdown files from a directory.

        Returns:
            List of dicts with content, source, source_url
        """
        documents = []
        md_dir = Path(directory)

        if not md_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        for md_file in md_dir.glob("**/*.md"):
            try:
                content = self.load_text_file(str(md_file))
                if content.strip():
                    documents.append({
                        "content": content,
                        "source": "markdown",
                        "source_url": str(md_file),
                        "metadata": {
                            "file_name": md_file.name,
                            "file_path": str(md_file),
                        },
                    })
                    logger.info(f"Loaded markdown: {md_file.name}")
            except Exception as e:
                logger.error(f"Error processing {md_file}: {str(e)}")

        return documents

    def load_text_files(self, directory: str) -> List[Dict]:
        """Load all text files from a directory"""
        documents = []
        txt_dir = Path(directory)

        if not txt_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        for txt_file in txt_dir.glob("**/*.txt"):
            try:
                content = self.load_text_file(str(txt_file))
                if content.strip():
                    documents.append({
                        "content": content,
                        "source": "text",
                        "source_url": str(txt_file),
                        "metadata": {
                            "file_name": txt_file.name,
                            "file_path": str(txt_file),
                        },
                    })
                    logger.info(f"Loaded text file: {txt_file.name}")
            except Exception as e:
                logger.error(f"Error processing {txt_file}: {str(e)}")

        return documents

    def load_pdf_files(self, directory: str) -> List[Dict]:
        """Load all PDF files from a directory and extract text"""
        documents = []
        pdf_dir = Path(directory)

        if not pdf_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        if not PDF_SUPPORT:
            logger.warning("pypdf not installed - skipping PDF files. Install with: pip install pypdf")
            return documents

        for pdf_file in pdf_dir.glob("**/*.pdf"):
            try:
                reader = PdfReader(str(pdf_file))
                text = ""

                # Extract text from all pages
                for page_num, page in enumerate(reader.pages):
                    try:
                        text += page.extract_text() + "\n"
                    except Exception as e:
                        logger.warning(f"Could not extract page {page_num} from {pdf_file.name}: {str(e)}")

                if text.strip():
                    documents.append({
                        "content": text,
                        "source": "pdf",
                        "source_url": str(pdf_file),
                        "metadata": {
                            "file_name": pdf_file.name,
                            "file_path": str(pdf_file),
                            "pages": len(reader.pages),
                        },
                    })
                    logger.info(f"Loaded PDF: {pdf_file.name} ({len(reader.pages)} pages)")
            except Exception as e:
                logger.error(f"Error processing {pdf_file}: {str(e)}")

        return documents

    def load_docx_files(self, directory: str) -> List[Dict]:
        """Load all DOCX files from a directory and extract text"""
        documents = []
        docx_dir = Path(directory)

        if not docx_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        if not DOCX_SUPPORT:
            logger.warning("python-docx not installed - skipping DOCX files. Install with: pip install python-docx")
            return documents

        for docx_file in docx_dir.glob("**/*.docx"):
            try:
                doc = Document(str(docx_file))
                text = "\n".join([para.text for para in doc.paragraphs])

                if text.strip():
                    documents.append({
                        "content": text,
                        "source": "docx",
                        "source_url": str(docx_file),
                        "metadata": {
                            "file_name": docx_file.name,
                            "file_path": str(docx_file),
                            "paragraphs": len(doc.paragraphs),
                        },
                    })
                    logger.info(f"Loaded DOCX: {docx_file.name}")
            except Exception as e:
                logger.error(f"Error processing {docx_file}: {str(e)}")

        return documents

    def load_json_files(self, directory: str) -> List[Dict]:
        """Load all JSON files from a directory"""
        documents = []
        json_dir = Path(directory)

        if not json_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        for json_file in json_dir.glob("**/*.json"):
            try:
                with open(json_file, "r", encoding="utf-8") as f:
                    data = json.load(f)

                # Convert JSON to readable text
                if isinstance(data, dict):
                    text = json.dumps(data, indent=2)
                elif isinstance(data, list):
                    text = "\n".join([json.dumps(item, indent=2) if isinstance(item, dict) else str(item) for item in data])
                else:
                    text = str(data)

                if text.strip():
                    documents.append({
                        "content": text,
                        "source": "json",
                        "source_url": str(json_file),
                        "metadata": {
                            "file_name": json_file.name,
                            "file_path": str(json_file),
                        },
                    })
                    logger.info(f"Loaded JSON: {json_file.name}")
            except Exception as e:
                logger.error(f"Error processing {json_file}: {str(e)}")

        return documents

    def load_csv_files(self, directory: str) -> List[Dict]:
        """Load all CSV files from a directory"""
        documents = []
        csv_dir = Path(directory)

        if not csv_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        for csv_file in csv_dir.glob("**/*.csv"):
            try:
                rows = []
                with open(csv_file, "r", encoding="utf-8") as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        rows.append(row)

                if rows:
                    # Format as readable text
                    text = "CSV Data:\n"
                    for i, row in enumerate(rows, 1):
                        text += f"Record {i}:\n"
                        for key, value in row.items():
                            text += f"  {key}: {value}\n"
                        text += "\n"

                    documents.append({
                        "content": text,
                        "source": "csv",
                        "source_url": str(csv_file),
                        "metadata": {
                            "file_name": csv_file.name,
                            "file_path": str(csv_file),
                            "rows": len(rows),
                        },
                    })
                    logger.info(f"Loaded CSV: {csv_file.name} ({len(rows)} rows)")
            except Exception as e:
                logger.error(f"Error processing {csv_file}: {str(e)}")

        return documents

    def load_excel_files(self, directory: str) -> List[Dict]:
        """Load all Excel files from a directory"""
        documents = []
        excel_dir = Path(directory)

        if not excel_dir.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        if not EXCEL_SUPPORT:
            logger.warning("openpyxl not installed - skipping Excel files. Install with: pip install openpyxl")
            return documents

        for excel_file in excel_dir.glob("**/*.xlsx"):
            try:
                wb = openpyxl.load_workbook(str(excel_file), data_only=True)
                text = ""

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text += f"\n=== Sheet: {sheet_name} ===\n"
                    for row in ws.iter_rows(values_only=True):
                        text += " | ".join([str(cell) if cell is not None else "" for cell in row]) + "\n"

                if text.strip():
                    documents.append({
                        "content": text,
                        "source": "excel",
                        "source_url": str(excel_file),
                        "metadata": {
                            "file_name": excel_file.name,
                            "file_path": str(excel_file),
                            "sheets": len(wb.sheetnames),
                        },
                    })
                    logger.info(f"Loaded Excel: {excel_file.name} ({len(wb.sheetnames)} sheets)")
            except Exception as e:
                logger.error(f"Error processing {excel_file}: {str(e)}")

        return documents

    def load_sample_documents(self) -> List[Dict]:
        """
        Load sample documents for testing.

        Returns sample Vivli-related content
        """
        return [
            {
                "content": """
How to Submit a Data Request

A data request is your formal application to access data available on Vivli.

Steps to submit:
1. Create an account or log in
2. Click "New Data Request"
3. Fill in your research details
4. Upload required documents (CV, institution letter)
5. Submit for review

Timeline: Most requests are reviewed within 2 weeks.

Eligibility:
- Must be affiliated with a research institution
- Must have institutional approval
- Must agree to data use terms
                """,
                "source": "sample",
                "source_url": "sample:data-request-guide",
                "metadata": {
                    "title": "Data Request Submission Guide",
                    "category": "FAQ",
                },
            },
            {
                "content": """
Form Check Process

After you submit your data request, it goes through a form check.

What happens during form check:
1. Vivli team validates all required fields
2. Checks document completeness
3. Verifies institutional affiliation
4. May request additional information

If issues are found:
- You'll receive feedback via the platform
- You can revise and resubmit
- Revised submissions are re-checked

Typical form check time: 3-5 business days
                """,
                "source": "sample",
                "source_url": "sample:form-check",
                "metadata": {
                    "title": "Form Check Process",
                    "category": "FAQ",
                },
            },
            {
                "content": """
Data Review Process

Once your form passes validation, it enters the review stage.

Review includes:
1. Scientific merit assessment
2. Data access appropriateness
3. Compliance with terms
4. Ethics review if needed

Possible outcomes:
- APPROVED: Access granted immediately
- APPROVED WITH CONDITIONS: Need to meet additional requirements
- NEED MORE INFO: Please provide clarification
- REJECTED: Reason provided

Review timeline: 2-4 weeks typically

You can contact us if review takes longer than 30 days.
                """,
                "source": "sample",
                "source_url": "sample:review-process",
                "metadata": {
                    "title": "Data Review Process",
                    "category": "FAQ",
                },
            },
        ]

    def load_from_organized_data(self) -> List[Dict]:
        """
        Load documents from your organized-data directory.

        Loads multiple file formats:
        - Markdown files (.md)
        - PDF files (.pdf)
        - Text files (.txt)
        - Word documents (.docx)
        - JSON files (.json)
        - CSV files (.csv)
        - Excel files (.xlsx)

        Returns:
            List of loaded documents
        """
        documents = []
        base = Path(self.base_path)

        if not base.exists():
            logger.warning(f"Base path does not exist: {base}")
            return documents

        # Load all file types from all subdirectories
        logger.info("Loading markdown files...")
        documents.extend(self.load_markdown_files(str(base)))

        logger.info("Loading PDF files...")
        documents.extend(self.load_pdf_files(str(base)))

        logger.info("Loading text files...")
        documents.extend(self.load_text_files(str(base)))

        logger.info("Loading DOCX files...")
        documents.extend(self.load_docx_files(str(base)))

        logger.info("Loading JSON files...")
        documents.extend(self.load_json_files(str(base)))

        logger.info("Loading CSV files...")
        documents.extend(self.load_csv_files(str(base)))

        logger.info("Loading Excel files...")
        documents.extend(self.load_excel_files(str(base)))

        logger.info(f"Total documents loaded: {len(documents)}")
        return documents
